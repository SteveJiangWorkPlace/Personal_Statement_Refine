import os
import streamlit as st
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import re
import time
from io import BytesIO
from PIL import Image

# ==========================================
# 🔴 网络代理配置
# 设置HTTP和HTTPS代理，用于确保应用能够通过代理访问Google Gemini API
# ==========================================
# os.environ["HTTP_PROXY"] = "http://127.0.0.1:7897"
# os.environ["HTTPS_PROXY"] = "http://127.0.0.1:7897"

# ==========================================
# 依赖库检测与初始化
# 检查是否安装了处理Word文档和PDF文件的库，并相应设置标志
# ==========================================
HAS_DOCX = False
HAS_PDF = False

try:
    # 尝试导入处理Word文档的库
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    HAS_DOCX = True
except ImportError:
    pass

try:
    # 尝试导入处理PDF文件的库
    import pypdf
    HAS_PDF = True
except ImportError:
    pass

# ==========================================
# 自定义UI样式函数
# 通过注入CSS来创建米色背景和宝蓝色按钮的自定义界面
# ==========================================
def apply_custom_css():
    st.markdown("""
    <style>
    /* 引入 Inter 字体 */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    /* 全局变量 - 定制配色 */
    :root {
        --primary-color: #3666FA; /* 宝蓝 RGB 54, 102, 250 */
        --bg-color: #FBF7EC;      /* 米色 RGB 251, 247, 236 */
        --text-color: #3666FA;    /* 字体颜色跟随主色 */
        --button-text: #FBF7EC;   /* 按钮内文字颜色 (米色) */
    }

    /* 基础重置 */
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        color: var(--text-color);
        background-color: var(--bg-color);
    }
    
    /* 隐藏 Streamlit 默认 Header 和 Footer */
    header {visibility: hidden;}
    footer {visibility: hidden;}

    /* 主容器背景优化 */
    .stApp {
        background-color: var(--bg-color);
    }

    /* 侧边栏优化 - 深色沉浸式 */
    [data-testid="stSidebar"] {
        background-color: #0f172a; 
        border-right: 1px solid #1e293b;
    }
    
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3, 
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] div {
        color: #e2e8f0 !important;
    }
    
    [data-testid="stSidebar"] hr {
        border-color: #334155 !important;
    }

    /* 标题样式 - 左对齐，大字体 */
    h1 {
        color: var(--text-color) !important;
        font-weight: 800 !important;
        font-size: 2.5rem !important;
        letter-spacing: -0.02em;
        margin-bottom: 2rem !important;
        text-align: left !important;
    }
    
    /* 小标题样式 */
    h2, h3 {
        color: var(--text-color) !important;
        font-weight: 600 !important;
        margin-top: 1rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* 普通文本和Label颜色 */
    p, label, .stMarkdown, .stText {
        color: var(--text-color) !important;
    }

    /* 输入框美化 */
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        border: none !important;
        border-radius: 8px !important;
        padding: 0.6rem 0.8rem !important;
        background-color: #ffffff !important;
        font-size: 0.95rem !important;
        color: #1e293b !important; /* 输入框内部文字深色 */
        transition: all 0.2s ease;
    }

    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(54, 102, 250, 0.1) !important;
    }

    /* 按钮美化 - 宝蓝背景，米色文字 */
    .stButton button {
        background-color: var(--primary-color) !important;
        color: var(--button-text) !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.5rem 1rem !important;
        font-weight: 500 !important;
        font-size: 0.9rem !important;
        box-shadow: 0 1px 2px rgba(54, 102, 250, 0.2) !important;
        transition: all 0.2s ease !important;
        min-width: 100px !important;
        min-height: 38px !important;
        height: auto;
        line-height: 1.4;
    }
    
    /* 强制按钮内所有元素颜色为米色 */
    .stButton button * {
        color: var(--button-text) !important;
    }

    .stButton button:hover {
        opacity: 0.9;
        transform: translateY(-1px);
    }
    
    /* 下载按钮 */
    .stDownloadButton button {
        background-color: var(--primary-color) !important;
        color: #FFFFFF !important; /* 修改为白色文字 */
        border: none !important;
    }
                
    /* 强制下载按钮内所有元素颜色一致 */
    .stDownloadButton button * {
        color: #FFFFFF !important; /* 确保按钮内所有元素都是白色 */            
    }
                
    .stDownloadButton button:hover {
        opacity: 0.9;
    }

    /* Expander 样式微调 - 增加字重以支持加粗效果 */
    .streamlit-expanderHeader {
        background-color: #ffffff !important;
        border: 1px solid rgba(54, 102, 250, 0.2) !important;
        border-radius: 8px !important;
        color: var(--text-color) !important;
        font-weight: 600 !important; /* 强制加粗 */
    }
    
    /* 文件上传区域 */
    [data-testid="stFileUploader"] {
        border: 1px dashed rgba(54, 102, 250, 0.4);
        background-color: #ffffff;
        border-radius: 8px;
        padding: 1rem;
        min-height: 150px; /* 确保与文本框高度一致 */
        display: flex;
        flex-direction: column;
        justify-content: center;
        width: 100% !important;
        box-sizing: border-box !important;
    }
    [data-testid="stFileUploader"]:hover {
        border-color: var(--primary-color);
        background-color: rgba(54, 102, 250, 0.05);
    }

    /* 布局间距调整 */
    .block-container {
        padding-top: 3rem !important;
        padding-bottom: 3rem !important;
        max-width: 1200px !important;
    }
    
    /* 分割线颜色 */
    hr {
        border-color: rgba(54, 102, 250, 0.2) !important;
    }
    
    /* 进度条颜色 */
    .stProgress > div > div > div > div {
        background-color: var(--primary-color) !important;
    }
    
    /* 添加高亮样式 */
    .highlight {
        background-color: #FFEB3B;
        font-weight: bold;
    }
    
    /* 统一文本框样式 */
    .stTextArea textarea {
        border: 1px solid rgba(54, 102, 250, 0.2) !important;
        border-radius: 8px !important;
        padding: 10px !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif !important;
        font-size: 16px !important;
        line-height: 1.5 !important;
        color: #333333 !important;
        background-color: #ffffff !important;
        min-height: 300px !important;  /* 最小高度，允许扩展 */
        width: 100% !important;
        box-sizing: border-box !important;
    }
    
    /* 预览容器样式 */
    .preview-container {
        border: 1px solid rgba(54, 102, 250, 0.2);
        border-radius: 8px;
        padding: 10px;
        background-color: #ffffff;
        height: 300px;
        overflow-y: auto;
        margin-top: 10px; /* 与文本区域对齐 */
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif;
        font-size: 16px;
        line-height: 1.5;
        color: #333;
    }
    
    /* 批注结果预览容器 */
    .annotation-result-container {
        border: 1px solid rgba(54, 102, 250, 0.2);
        border-radius: 8px;
        padding: 10px;
        background-color: #ffffff;
        height: 300px;
        overflow-y: auto;
        margin-top: 10px;
        margin-bottom: 20px;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif;
        font-size: 16px;
        line-height: 1.5;
        color: #333;
    }
    
    /* 预览标题样式 */
    .preview-title {
        color: #3666FA;
        margin-bottom: 10px;
        font-weight: bold;
        font-size: 16px;
    }
    
    /* 预览文本样式 */
    .preview-text {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', sans-serif;
        font-size: 16px;
        line-height: 1.5;
        color: #333;
    }
    
    /* 统一信息框样式 */
    .stAlert {
        border-radius: 8px !important;
    }
    
    /* 调整列间距 */
    [data-testid="column"] {
        padding: 0 10px !important;
    }
    
    /* 确保预览区域与文本框对齐 */
    .preview-wrapper {
        height: 100%;
        display: flex;
        flex-direction: column;
    }
    
    /* 修改部分高亮显示 */
    .modified-text {
        background-color: #FFEB3B;
        font-weight: bold;
    }
    
    /* 确保上传文件区域和文本框顶端对齐 */
    .top-align-container {
        display: flex;
        align-items: flex-start;
    }
    
    /* 移除上传文件区域的上边距 */
    .top-align-container [data-testid="stFileUploader"] {
        margin-top: 0 !important;
    }
    
    /* 移除文本区域的上边距 */
    .top-align-container .stTextArea {
        margin-top: 0 !important;
    }

    /* 流式显示光标效果 */
    .streaming-cursor::after {
        content: "▌";
        animation: blink 1s infinite;
        color: var(--primary-color);
    }
    @keyframes blink {
        0%, 50% { opacity: 1; }
        51%, 100% { opacity: 0; }
    }

    /* 统一按钮样式 */
    .uniform-button button {
        min-width: 100px !important;
        min-height: 38px !important;
        width: 100% !important;
        box-sizing: border-box !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
    }

    /* 确认内容按钮特殊样式 */
    button[data-testid*="confirm_p_"] {
        background-color: white !important;
        color: #3666FA !important;
        border: 2px solid #3666FA !important;
    }

    button[data-testid*="confirm_p_"]:hover {
        opacity: 0.8;
    }
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 页面配置与会话状态初始化
# 设置页面标题、布局和初始化所有必要的会话状态变量
# ==========================================
st.set_page_config(page_title="个人陈述修改", layout="wide")

# 应用自定义UI样式
apply_custom_css()

# 调试模式标志
DEBUG_MODE = True

# 日志系统
import logging
from datetime import datetime

def setup_logging():
    """设置日志系统"""
    logger = logging.getLogger('psr_debug')
    logger.setLevel(logging.DEBUG)

    # 避免重复添加handler
    if not logger.handlers:
        # 文件handler
        fh = logging.FileHandler('psr_debug.log', encoding='utf-8')
        fh.setLevel(logging.DEBUG)

        # 控制台handler（仅当DEBUG_MODE开启时）
        if DEBUG_MODE:
            ch = logging.StreamHandler()
            ch.setLevel(logging.DEBUG)

        # 格式化
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        fh.setFormatter(formatter)
        if DEBUG_MODE:
            ch.setFormatter(formatter)

        logger.addHandler(fh)
        if DEBUG_MODE:
            logger.addHandler(ch)

    return logger

logger = setup_logging()

def log_session_state_summary():
    """记录session state的摘要信息"""
    logger.info("=== Session State 摘要 ===")
    logger.info(f"sections_data长度: {len(st.session_state.get('sections_data', []))}")
    logger.info(f"confirmed_paragraphs: {st.session_state.get('confirmed_paragraphs', set())}")
    logger.info(f"confirmed_contents keys: {list(st.session_state.get('confirmed_contents', {}).keys())}")

    # 检查confirmed_contents中的实际内容
    confirmed_contents = st.session_state.get('confirmed_contents', {})
    for idx, content in confirmed_contents.items():
        logger.info(f"confirmed_contents[{idx}] 长度: {len(content) if content else 0}")
        if content and len(content) < 100:
            logger.debug(f"confirmed_contents[{idx}] 内容: {content}")

    logger.info(f"final_preview_text长度: {len(st.session_state.get('final_preview_text', ''))}")
    logger.info(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
    logger.info(f"generation_complete: {st.session_state.get('generation_complete', False)}")
    logger.info(f"show_sections: {st.session_state.get('show_sections', False)}")
    logger.info("=== Session State 摘要结束 ===")

# 在脚本开始时记录session state
log_session_state_summary()

# 初始化所有会话状态变量，用于在页面重新加载时保持数据
if 'ps_content' not in st.session_state: st.session_state['ps_content'] = ""  # 原始PS内容
if 'curr_content' not in st.session_state: st.session_state['curr_content'] = ""  # 课程内容
if 'strategy_content' not in st.session_state: st.session_state['strategy_content'] = ""  # 策略内容
if 'sections_data' not in st.session_state: st.session_state['sections_data'] = []  # 段落数据
if 'translation_results' not in st.session_state: st.session_state['translation_results'] = {}  # 翻译结果
if 'edited_translations' not in st.session_state: st.session_state['edited_translations'] = {}  # 编辑后的翻译
if 'refine_results' not in st.session_state: st.session_state['refine_results'] = {}  # 修改结果
if 'preview_results' not in st.session_state: st.session_state['preview_results'] = {}  # 预览结果
if 'generation_complete' not in st.session_state: st.session_state['generation_complete'] = False  # 生成完成标志
if 'full_response' not in st.session_state: st.session_state['full_response'] = ""  # 完整响应
if 'show_sections' not in st.session_state: st.session_state['show_sections'] = False  # 显示段落标志
if 'annotation_processing' not in st.session_state: st.session_state['annotation_processing'] = {}  # 批注处理状态
if 'annotation_results' not in st.session_state: st.session_state['annotation_results'] = {}  # 批注处理结果
if 'original_texts' not in st.session_state: st.session_state['original_texts'] = {}  # 原始文本，用于比较
if 'final_preview_text' not in st.session_state: st.session_state['final_preview_text'] = ""  # 最终预览文本
if 'final_preview_text_cleaned' not in st.session_state: st.session_state['final_preview_text_cleaned'] = ""  # 清理后的最终预览文本
if 'confirmed_paragraphs' not in st.session_state: st.session_state['confirmed_paragraphs'] = set()  # 已确认段落的索引
if 'confirmed_contents' not in st.session_state: st.session_state['confirmed_contents'] = {}  # 已确认段落的内容

# 从Streamlit secrets获取Google API Key
api_key = st.secrets.get("GOOGLE_API_KEY")
if api_key:
    os.environ["GOOGLE_API_KEY"] = api_key
    genai.configure(api_key=api_key)
else:
    pass  # 错误信息在侧边栏中显示

# 侧边栏设置
with st.sidebar:
    st.markdown("### 设置")
    if api_key:
        st.success("✅ API Key 已从 Secrets 加载")
    else:
        st.sidebar.error("❌ API Key 未配置")
    st.divider()
    # 显示已生成段落的数量
    if st.session_state['sections_data']:
        st.success(f"当前已生成 {len(st.session_state['sections_data'])} 个段落")

    st.divider()
    st.markdown("### 诊断信息")

    # 显示最终预览状态
    if st.session_state['final_preview_text']:
        st.info(f"最终预览文本长度: {len(st.session_state['final_preview_text'])} 字符")
        # 显示前100字符预览
        preview_text = st.session_state['final_preview_text']
        if len(preview_text) > 100:
            st.text_area("final_preview_text预览", value=preview_text[:500], height=150, key="sidebar_preview", disabled=True)
        else:
            st.text_area("final_preview_text预览", value=preview_text, height=150, key="sidebar_preview", disabled=True)

        if st.session_state.get('final_preview_text_cleaned'):
            st.info(f"清理版本长度: {len(st.session_state['final_preview_text_cleaned'])} 字符")
            cleaned_preview = st.session_state['final_preview_text_cleaned']
            if len(cleaned_preview) > 100:
                st.text_area("final_preview_text_cleaned预览", value=cleaned_preview[:500], height=150, key="sidebar_cleaned_preview", disabled=True)
            else:
                st.text_area("final_preview_text_cleaned预览", value=cleaned_preview, height=150, key="sidebar_cleaned_preview", disabled=True)
    else:
        st.warning("最终预览文本为空")

    # 显示已确认段落状态
    confirmed_count = len(st.session_state['confirmed_paragraphs'])
    total_paragraphs = len(st.session_state['sections_data'])
    if total_paragraphs > 0:
        st.info(f"已确认段落: {confirmed_count}/{total_paragraphs}")
        if confirmed_count > 0:
            confirmed_indices = sorted(list(st.session_state['confirmed_paragraphs']))
            st.info(f"已确认段落索引: {confirmed_indices}")

    # 检查confirmed_contents
    if st.session_state['confirmed_contents']:
        st.info(f"confirmed_contents中有 {len(st.session_state['confirmed_contents'])} 个段落")
        for idx, content in st.session_state['confirmed_contents'].items():
            content_len = len(content) if content else 0
            st.info(f"段落 {idx}: {content_len} 字符")
    else:
        st.warning("confirmed_contents为空")

    # 诊断按钮
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📋 输出详细诊断日志", key="diagnostic_btn"):
            logger.info(f"=== 详细诊断日志 ===")
            logger.info(f"final_preview_text长度: {len(st.session_state['final_preview_text'])}")
            logger.info(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
            logger.info(f"final_preview_text_display session state存在: {'final_preview_text_display' in st.session_state}")
            logger.info(f"sections_data长度: {len(st.session_state['sections_data'])}")
            logger.info(f"confirmed_paragraphs: {st.session_state['confirmed_paragraphs']}")
            logger.info(f"confirmed_contents keys: {list(st.session_state['confirmed_contents'].keys())}")

            # 检查每个confirmed_contents的内容
            for idx, content in st.session_state['confirmed_contents'].items():
                logger.info(f"confirmed_contents[{idx}]长度: {len(content) if content else 0}")
                if content and len(content) < 500:
                    logger.info(f"confirmed_contents[{idx}]内容前200字符: {content[:200]}")

            # 检查final_preview_text_display
            if 'final_preview_text_display' in st.session_state:
                display_val = st.session_state['final_preview_text_display']
                logger.info(f"final_preview_text_display长度: {len(display_val) if display_val else 0}")
                if display_val and len(display_val) < 500:
                    logger.info(f"final_preview_text_display前200字符: {display_val[:200]}")

            # 检查display_text的计算
            cleaned_text = st.session_state.get('final_preview_text_cleaned', '')
            if cleaned_text and cleaned_text.strip():
                display_text = cleaned_text
                logger.info(f"display_text使用final_preview_text_cleaned，长度: {len(display_text)}")
            else:
                display_text = st.session_state['final_preview_text']
                logger.info(f"display_text使用final_preview_text，长度: {len(display_text)}")
            logger.info(f"display_text前200字符: {display_text[:200] if display_text else '空'}")

            logger.info(f"=== 诊断日志结束 ===")
            st.success("详细诊断日志已输出到日志文件")

    with col2:
        if st.button("🔄 强制重建预览", key="rebuild_preview_btn"):
            logger.info(f"=== 强制重建预览 ===")
            rebuilt_text = rebuild_final_preview()
            logger.info(f"重建结果长度: {len(rebuilt_text)}")
            if rebuilt_text and rebuilt_text.strip():
                st.session_state['final_preview_text'] = rebuilt_text
                logger.info(f"final_preview_text已更新，长度: {len(rebuilt_text)}")
                st.success(f"预览已重建，长度: {len(rebuilt_text)} 字符")
                # 清除清理版本
                st.session_state['final_preview_text_cleaned'] = ''
                logger.info("已清除final_preview_text_cleaned")
                st.rerun()
            else:
                logger.warning(f"重建结果为空")
                st.error("重建失败，结果为空")

# 设置默认使用的模型
model_name = "gemini-2.5-pro"

# 安全设置，用于交互式API调用
safety_settings_interactive = {
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# ==========================================
# 工具函数
# 包含各种辅助功能，如文件处理、文本清理和格式转换
# ==========================================

# 从上传的文件中提取文本内容
def extract_text_from_file(uploaded_file):
    """从上传的文件中提取文本，支持DOCX、PDF和TXT格式"""
    if not uploaded_file: return ""
    file_type = uploaded_file.name.split('.')[-1].lower()
    text = ""
    try:
        if file_type == 'docx' and HAS_DOCX:
            doc = Document(uploaded_file)
            for para in doc.paragraphs: text += para.text + "\n"
        elif file_type == 'pdf' and HAS_PDF:
            reader = pypdf.PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text() + "\n"
        elif file_type == 'txt':
            text = uploaded_file.getvalue().decode("utf-8")
    except Exception as e:
        return f"[读取文件出错: {e}]"
    return text

# 清除文本中的星号
def clean_asterisks(text):
    """移除文本中的所有星号字符"""
    if not text: return ""
    return text.replace("*", "")

# 移除Markdown加粗标记
def remove_markdown_bold(text):
    """移除文本中的Markdown加粗标记（**）"""
    return text.replace("**", "")

# 过滤AI生成内容中的问候语
def filter_ai_greeting(text):
    """移除AI生成内容开头的常见问候语和介绍语"""
    greeting_patterns = [
        r'^好的，作为.*?顾问.*?\n+',
        r'^作为.*?顾问.*?\n+',
        r'^我将.*?分析.*?\n+',
        r'^下面我将.*?\n+',
        r'^我会.*?帮助您.*?\n+',
        r'^让我.*?为您.*?\n+'
    ]
    
    for pattern in greeting_patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)
    
    return text

# 创建带有格式的Word文档
def create_docx_smart(text_content, major_name=""):
    """创建格式化的Word文档，包括页眉、字体设置和加粗高亮"""
    if not HAS_DOCX: return None
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加页眉
    header_text = f"Personal Statement - {major_name}" if major_name else "Personal Statement"
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = header_text
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置页眉文本格式
    header_run = header_para.runs[0]
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(11)

    # 设置正文默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 处理正文内容，保留加粗格式
    lines = text_content.split('\n')
    for line in lines:
        if not line.strip(): continue
        clean_line = line.replace('[[LOGIC]]', '').replace('[[DRAFT]]', '')
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', clean_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2]
                run = p.add_run(clean_text)
                run.bold = True
            else:
                p.add_run(part)
    
    # 将文档保存到内存缓冲区
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 生成HTML预览，高亮显示加粗部分
def generate_preview_html(text_with_markdown):
    """将Markdown格式的文本转换为HTML预览，高亮显示加粗部分"""
    # 替换markdown加粗语法为HTML span标签
    html_text = re.sub(r'\*\*(.*?)\*\*', r'<span style="background-color: #FFEB3B; font-weight: bold;">\1</span>', text_with_markdown)
    
    # 添加HTML样式，确保与文本框样式一致
    styled_html = f"""
    <div class="preview-container">
        <div class="preview-text">
            {html_text}
        </div>
    </div>
    """
    return styled_html

# 新增函数：比较文本并高亮差异部分
def highlight_differences(original_text, new_text):
    """比较原始文本和新文本，高亮显示差异部分"""
    # 这是一个简化的实现，实际上需要更复杂的文本差异比较算法
    # 这里我们使用一个简单的方法：将新文本中的每个句子与原文本比较
    
    # 如果原文本为空，则将整个新文本高亮显示
    if not original_text:
        return f"<span class='modified-text'>{new_text}</span>"
    
    # 将文本分割成句子
    def split_into_sentences(text):
        # 使用正则表达式分割句子，考虑各种标点符号
        return re.split(r'([.!?。！？\n]+)', text)
    
    orig_sentences = split_into_sentences(original_text)
    new_sentences = split_into_sentences(new_text)
    
    # 合并相邻的分割结果
    orig_sentences_merged = []
    for i in range(0, len(orig_sentences)-1, 2):
        if i+1 < len(orig_sentences):
            orig_sentences_merged.append(orig_sentences[i] + orig_sentences[i+1])
        else:
            orig_sentences_merged.append(orig_sentences[i])
    
    new_sentences_merged = []
    for i in range(0, len(new_sentences)-1, 2):
        if i+1 < len(new_sentences):
            new_sentences_merged.append(new_sentences[i] + new_sentences[i+1])
        else:
            new_sentences_merged.append(new_sentences[i])
    
    # 标记每个新句子是否存在于原文本中
    result = []
    for sentence in new_sentences_merged:
        if sentence.strip() and sentence.strip() not in original_text:
            result.append(f"<span class='modified-text'>{sentence}</span>")
        else:
            result.append(sentence)
    
    # 合并结果
    return "".join(result)

# 检测文本是否包含中文
def contains_chinese(text):
    """检测文本中是否包含中文字符"""
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

# 检测文本是否包含批注标记
def contains_annotation(text):
    """检测文本是否包含【】或[]形式的批注标记"""
    return ('【' in text and '】' in text) or ('[' in text and ']' in text)

# 从AI修改思路中提取段落主题
def extract_paragraph_topic(logic_text):
    """从AI修改思路中提取段落主题"""
    if not logic_text:
        return "未识别"

    # 尝试从常见模式中提取
    patterns = [
        r"本段功能识别：\[(.+?)\]",
        r"功能：(.+?)(?:\n|$)",
        r"主题：(.+?)(?:\n|$)"
    ]

    for pattern in patterns:
        match = re.search(pattern, logic_text)
        if match:
            return match.group(1).strip()

    # 根据关键词推断
    keywords = {
        "动机": ["动机", "兴趣", "inspiration", "motivation"],
        "学术背景": ["学术", "学习", "课程", "academic"],
        "研究经历": ["研究", "项目", "实验", "research"],
        "工作经历": ["工作", "实习", "职业", "work"],
        "职业规划": ["规划", "目标", "未来", "career"],
        "择校理由": ["学校", "课程", "专业", "why school"]
    }

    for topic, key_list in keywords.items():
        if any(key in logic_text.lower() for key in key_list):
            return topic

    return "段落内容"

# 重建最终预览文本
def rebuild_final_preview():
    """按段落顺序重建最终预览文本"""
    logger.info(f"=== 开始重建最终预览 ===")
    logger.info(f"sections_data长度: {len(st.session_state.get('sections_data', []))}")
    logger.info(f"confirmed_paragraphs: {st.session_state.get('confirmed_paragraphs', set())}")
    logger.info(f"confirmed_contents keys: {list(st.session_state.get('confirmed_contents', {}).keys())}")

    if not st.session_state['sections_data']:
        logger.warning("没有段落数据")
        if DEBUG_MODE:
            st.warning("没有段落数据")  # 调试信息
        return ""

    # 确保confirmed_paragraphs是list类型
    if isinstance(st.session_state['confirmed_paragraphs'], set):
        confirmed_indices = sorted(list(st.session_state['confirmed_paragraphs']))
    else:
        confirmed_indices = sorted(st.session_state['confirmed_paragraphs'])

    logger.info(f"confirmed_indices: {confirmed_indices}")

    if not confirmed_indices:
        logger.warning(f"已确认段落为空: {st.session_state['confirmed_paragraphs']}")
        if DEBUG_MODE:
            st.warning(f"已确认段落为空: {st.session_state['confirmed_paragraphs']}")  # 调试信息
        return ""

    paragraphs = []
    for idx in confirmed_indices:
        if idx < len(st.session_state['sections_data']):
            # 调试输出
            has_content = idx in st.session_state['confirmed_contents']
            logger.info(f"处理段落 {idx}, confirmed_contents中有: {has_content}")

            if DEBUG_MODE:
                st.info(f"处理段落 {idx}, confirmed_contents中有: {has_content}")

            if idx in st.session_state['confirmed_contents']:
                current_text = st.session_state['confirmed_contents'][idx]
                logger.info(f"段落 {idx} 从confirmed_contents获取内容，长度: {len(current_text) if current_text else 0}")
                logger.debug(f"段落 {idx} 内容前100字符: {current_text[:100] if current_text else '空'}")

                # 如果confirmed_contents中的内容为空，尝试从其他地方获取
                if not current_text or not current_text.strip():
                    logger.warning(f"段落 {idx} confirmed_contents中的内容为空，尝试从其他地方获取")
                    textarea_key = f"draft_p_{idx}"
                    if textarea_key in st.session_state:
                        fallback_text = st.session_state[textarea_key]
                        if fallback_text and fallback_text.strip():
                            current_text = fallback_text
                            logger.info(f"段落 {idx} 从textarea获取替代内容，长度: {len(current_text)}")
                    else:
                        # 最后回退到段落原始内容
                        draft_key = f"para_{idx}"
                        fallback_text = st.session_state['refine_results'].get(draft_key, st.session_state['sections_data'][idx]['draft'])
                        if fallback_text and fallback_text.strip():
                            current_text = fallback_text
                            logger.info(f"段落 {idx} 回退到原始内容，长度: {len(current_text)}")
            else:
                # 如果没有保存的内容，尝试从文本框获取最新内容
                textarea_key = f"draft_p_{idx}"
                if textarea_key in st.session_state:
                    current_text = st.session_state[textarea_key]
                    logger.info(f"段落 {idx} 从textarea获取内容，长度: {len(current_text) if current_text else 0}")
                else:
                    # 最后回退到段落原始内容
                    draft_key = f"para_{idx}"
                    current_text = st.session_state['refine_results'].get(draft_key, st.session_state['sections_data'][idx]['draft'])
                    logger.info(f"段落 {idx} 回退到原始内容，长度: {len(current_text) if current_text else 0}")

            if current_text and current_text.strip():
                paragraphs.append(current_text)
                logger.info(f"段落 {idx} 已添加到paragraphs列表，长度: {len(current_text)}")
            else:
                logger.warning(f"段落 {idx} 内容为空或仅空白字符")

    result = "\n\n".join(paragraphs)
    logger.info(f"最终结果长度: {len(result)}")
    logger.debug(f"最终结果前200字符: {result[:200] if result else '空'}")

    if DEBUG_MODE:
        st.info(f"重建结果长度: {len(result)}")  # 调试信息

    logger.info(f"=== 重建完成 ===")
    return result

# ==========================================
# Prompt构建函数
# 为不同任务创建专门的提示词，如分析、修改和翻译
# ==========================================

# 构建初始分析提示词
def build_analysis_prompt(school, major, old_text, new_course_text, has_images, strategy_text):
    """构建用于初始分析和生成中英混合文本的提示词"""
    # 如果上传了图片，添加相关指示
    image_instruction = "我同时也上传了课程设置的截图，请务必结合截图内容。" if has_images else ""
    
    # 如果提供了策略文本，添加到提示中
    custom_strategy_instruction = ""
    if strategy_text and strategy_text.strip():
        custom_strategy_instruction = f"""
        【用户特别指令 (优先级最高)】
        {strategy_text}
        """
    
    # 返回完整的提示词
    return f"""
    你是一位专业的留学文书顾问。
    【任务目标】将用户的【旧个人陈述】适配到新的申请目标：**{school}** 的 **{major}** 专业。
    {custom_strategy_instruction}
    【输入材料】
    1. 旧 PS 内容：
    {old_text}
    2. 新项目课程信息：
    {new_course_text}
    {image_instruction}
    
    【核心修改逻辑 (必须严格执行)】
    1. **结构与顺序 (尊重原文)**：
       - 请**顺应旧文书原本的段落结构和逻辑顺序**进行输出，不要强行打乱或重组。
       - **关键要求**：在处理每一段时，你必须在 `[[LOGIC]]` 中明确识别出**这一段的功能**。
    
    2. **针对"课程设置/择校理由"段落 (智能识别并深度重写)**：
       - 当你处理到**涉及学校、课程、Why School**的段落时，必须**完全重写**。
       - **筛选逻辑**：排除通用课程，只选与学生背景结合紧密的核心课。
       - **深度与具体化**：必须深入引用该课程模块中的**关键概念 (Key Concepts)** 或 **具体方法学**。

    3. **针对其他段落 (全篇适配与优化)**：
       - **范围覆盖**：开头动机、学习/实践经历、职业规划。
       - **适配新专业**：检查内容是否符合新专业逻辑。

    【⚠️⚠️⚠️ 绝对强制执行规则 (ABSOLUTE MANDATORY RULES) ⚠️⚠️⚠️】
    在生成 `[[DRAFT]]` 时，必须严格执行以下"中英混合"逻辑，这是最高优先级指令：
    1. **Unchanged Parts (未修改部分)**: MUST remain in **Original English**. Do NOT translate them into Chinese. 未修改部分必须保留原始英文。
    2. **Modified/New Parts (修改/新增部分)**: MUST be written in **CHINESE (中文)** directly without any brackets or parentheses. 所有修改或新增的部分必须直接用中文写出，不要用任何符号包裹。
       - Example: Original English text... 这里插入一句关于课程 A 的具体分析，强调它如何提升我的数据挖掘能力... more original English text.
    3. **Rewrite Sections (重写段落)**: If a whole paragraph (like Why School) is rewritten, output it **entirely in Chinese** without any brackets. 如果整段重写（如Why School段落），必须将整段内容直接用中文写出。
       - Example: 整段重写的内容...
    
    【⚠️ 严格禁止】
    1. 不要在输出开头添加任何问候语或介绍语，如"作为一名专业的留学文书顾问..."
    2. 直接从第一段内容开始输出，不要有任何前言或开场白
    3. 所有修改过的内容必须用中文表达，不要直接输出英文修改
    4. 不要用英文输出任何修改内容，所有修改必须是中文
    5. 不要使用任何符号（如方括号[]、圆括号()等）来包裹中文内容，直接输出中文即可

    【输出格式示例】
    ===SECTION===
    [[LOGIC]]
    本段功能识别：[例如：学术背景]
    这里用中文解释修改思路...
    [[DRAFT]]
    Original English sentence here. 这里插入一句补充说明，强调量化能力. Another original English sentence.
    ===SECTION===
    ...

    请开始输出：
    """

# 构建修改提示词 - 修改后确保直接替换原文本
def build_refine_prompt(text_with_instructions, has_chinese):
    """构建用于根据批注修改文本的提示词，根据文本是否包含中文决定输出语言"""
    # 根据文本是否包含中文决定输出语言
    output_language = "CHINESE" if has_chinese else "ENGLISH"
    
    return f"""
    You are an expert editor. The user has provided a draft text below, but they have inserted **modification instructions** inside brackets `【...】` or `[...]`.
    **Your Task:**
    1. Read the text carefully.
    2. Identify the instructions inside `【】` or `[]` (e.g., "【把这段语气改得更自信一点】", "[make this more professional]").
    3. **Execute** these instructions to rewrite the text.
    4. **Remove** the instruction markers and the instruction text itself from the final output.
    5. Keep the rest of the text that was not targeted by instructions unchanged.
    6. Ensure the final output is smooth and coherent.
    
    **IMPORTANT OUTPUT LANGUAGE RULE:**
    - The text contains Chinese: {has_chinese}
    - Your output MUST be in {output_language}. 
    - If the input contains Chinese text, keep using Chinese in your output.
    - If the input is entirely in English, respond in English.
    
    **Input Text:**
    {text_with_instructions}
    **Output:**
    Output ONLY the refined text (no explanations).
    """

# 修改翻译prompt，明确指示将中文翻译为英文，确保输出纯英文且无Markdown符号
def build_translate_prompt(hybrid_text, style="US"):
    """构建用于将中英混合文本翻译为纯英文的提示词，支持美式和英式拼写，遵循专业写作规范"""
    # 根据指定风格设置拼写规则
    spelling_rule = "American Spelling (Color, Honor, Analyze)" if style == "US" else "British Spelling (Colour, Honour, Analyse)"

    return f"""
    You are an expert Admissions Essay Translator.
    Task: Translate the hybrid Chinese-English paragraph into professional English.
    Spelling Convention: {spelling_rule}.
    Input (Hybrid Draft):
    {hybrid_text}
    CRITICAL RULES (MUST FOLLOW)
    1. **TRANSLATION EXECUTION**:
       - **MUST translate ALL Chinese text** into professional English following the rules below.
       - Any text inside brackets like `(...)` or `【...】` must be translated to English.
       - Merge translations smoothly with the existing English text.
       - **DO NOT use any Markdown formatting symbols** (no asterisks, bold, etc.)
       - Output clean text without any formatting marks.
       - Output ONLY the final English paragraph.
    2. **BANNED VOCABULARY (DO NOT USE)**:
       - master / mastery
       - my goal is to
       - permit
       - deep comprehension
       - look forward to
       - address
       - command
       - drawn to / draw
       - privilege
       - testament
       - commitment
       - tenure
       - thereby / thereby doing
       - cultivate
       - Building on this / Building on this foundation
       - intend to
       - demonstrate (use sparingly, avoid frequent appearance)
    3. **PROHIBITED STRUCTURES (ABSOLUTELY FORBIDDEN)**:
       - **Adverbs**: Do not use adverbs (including adverbs as logical connectors).
       - **-ing forms as nouns**: Avoid using -ing forms as nouns (gerunds as subjects/objects).
       - **Adverb + verb/adjective structures**: Avoid combinations like "significantly improve" or "deeply understand".
       - **Main clause + , + -ing participial phrases**: Avoid structures like "I completed the project, demonstrating my skills".
    4. **SENTENCE STRUCTURE REQUIREMENTS**:
       - **Use subordinate clauses** to enhance logical connections. For example: "...which in turn leads to..." instead of "...this [verb]..."
       - **Use semicolons (;)** to connect complete but conceptually related sentences, not periods.
       - Ensure logical coherence and smooth flow.
    5. **PUNCTUATION STANDARDS**:
       - **Quotation marks**: Do NOT place commas or periods inside quotation marks. Place punctuation OUTSIDE quotation marks.
       - Example: Use "example", not "example,".
    6. **PROFESSIONAL WRITING STANDARDS**:
       - Use precise, professional terminology.
       - Avoid colloquial expressions.
       - Maintain formal academic tone appropriate for personal statements.
    7. **ORIGINAL ENGLISH PRESERVATION**:
       - Keep original English parts unchanged.
       - Apply all rules above only to newly translated parts (from Chinese to English).
    """

# 修改英文精修提示词，确保输出纯英文且无Markdown符号
def build_english_refine_prompt(text_with_instructions):
    """构建用于英文精修阶段的提示词，确保输出纯英文，遵循专业写作规范"""
    return f"""
    You are an expert academic editor specializing in personal statements for graduate school applications.

    **Your Task:**
    1. Read the English text carefully.
    2. Identify the instructions inside `【】` or `[]` (e.g., "[make this more professional]", "【improve this sentence】").
    3. **Execute** these instructions to improve the text.
    4. **Remove** the instruction markers and the instruction text itself from the final output.
    5. Keep the rest of the text that was not targeted by instructions unchanged.
    6. Ensure the final output is smooth, coherent, and maintains a professional academic tone.

    **CRITICAL RULES (MUST FOLLOW):**
    1. **OUTPUT FORMAT**:
       - Output MUST be in ENGLISH only.
       - **DO NOT use any Markdown formatting symbols** (no asterisks, bold, etc.)
       - Output clean text without any formatting marks.

    2. **BANNED VOCABULARY (DO NOT USE)**:
       - master / mastery
       - my goal is to
       - permit
       - deep comprehension
       - look forward to
       - address
       - command
       - drawn to / draw
       - privilege
       - testament
       - commitment
       - tenure
       - thereby / thereby doing
       - cultivate
       - Building on this / Building on this foundation
       - intend to
       - demonstrate (use sparingly, avoid frequent appearance)

    3. **PROHIBITED STRUCTURES (ABSOLUTELY FORBIDDEN)**:
       - **Adverbs**: Do not use adverbs (including adverbs as logical connectors).
       - **-ing forms as nouns**: Avoid using -ing forms as nouns (gerunds as subjects/objects).
       - **Adverb + verb/adjective structures**: Avoid combinations like "significantly improve" or "deeply understand".
       - **Main clause + , + -ing participial phrases**: Avoid structures like "I completed the project, demonstrating my skills".

    4. **SENTENCE STRUCTURE REQUIREMENTS**:
       - **Use subordinate clauses** to enhance logical connections. For example: "...which in turn leads to..." instead of "...this [verb]..."
       - **Use semicolons (;)** to connect complete but conceptually related sentences, not periods.
       - Ensure logical coherence and smooth flow.

    5. **PUNCTUATION STANDARDS**:
       - **Quotation marks**: Do NOT place commas or periods inside quotation marks. Place punctuation OUTSIDE quotation marks.
       - Example: Use "example", not "example,".

    6. **PROFESSIONAL WRITING STANDARDS**:
       - Use precise, professional terminology.
       - Avoid colloquial expressions.
       - Maintain formal academic tone appropriate for personal statements.
       - Maintain the original meaning and intent of the text.

    **Input Text:**
    {text_with_instructions}

    **Output:**
    Output ONLY the refined English text (no explanations, no formatting marks).
    """

# 构建去除AI写作高频词汇的提示词
def build_remove_ai_vocab_prompt(text):
    """构建用于去除AI写作高频词汇和句式的提示词"""
    return f"""
你是一位专业的英文写作编辑，任务是去除个人陈述中的AI写作高频词汇和句式，使文本更加自然、个性化。

**绝对禁用的AI词汇和句式（黑名单）：**
A. 滥用的词汇和短语：
动词：
address (问题)
cultivate
Demonstrate（非严格禁用，需要少用，不要多次重复出现）
draw (特指 "draw from experience" 这类用法)
master
permit
leverage, utilize
名词和名词短语：
command (of a skill)
commitment
comprehension (尤其是 deep comprehension)
Master/mastery
privilege
tenure
testament
陈腐短语：
Building on this... / Building on this foundation
drawn to
look forward to
my goal is to
intend to
B. 滥用的结构和比喻：
副词+动词/形容词结构：避免过度使用"显著提升"、"深入理解"这类组合。
公式化因果：禁用 By doing X, I was able to Y 和 ...thereby doing... 的句式。
陈腐的比喻：
"旅程"隐喻 (e.g., academic/career journey)
"工具箱"隐喻 (e.g., skill set/toolkit)
"交汇点"逻辑 (e.g., the intersection of X and Y)

C. **Sentence Structure Variety (Balanced Rule)**: AI models often overuse the "comma + verb-ing" structure (e.g., ", revealing trends"). Do not strictly ban it, as it is valid in academic English, but **use it sparingly** to avoid a repetitive "AI tone." Instead, prioritize variety by using relative clauses (e.g., ", which revealed..."), coordination (e.g., "and revealed..."), or starting new sentences where appropriate for better flow.

**重要规则：**
7. **IMPORTANT - Remove Markdown**: Remove all Markdown formatting symbols like asterisks (*), double asterisks (**), underscores (_), etc. from the output. Provide clean text without any Markdown formatting.
8. **Punctuation with Quotation Marks**: For general text (not formal citations), always place commas, periods, and other punctuation marks OUTSIDE of quotation marks, not inside. For example, use "example", not "example,". For formal citations, maintain the original citation style's punctuation rules.

**你的任务：**
1. 仔细阅读以下文本。
2. 识别并移除所有黑名单中的词汇和短语。
3. 改写包含禁用句式的句子，保持原意但使用更自然的表达。
4. 去除任何陈腐的比喻和公式化结构。
5. 使文本更加个性化、生动，避免AI生成的痕迹。
6. 保持文本的专业性和学术性。
7. **不要添加任何额外解释**，只输出修改后的文本。

**重要规则：**
- 只修改确实属于黑名单的内容，如果没有问题，不要随意修改。
- 保留文本的原始含义和逻辑。
- 输出语言与输入语言一致（英文输入则英文输出，中文输入则中文输出）。

**输入文本：**
{text}

**输出：**
只输出修改后的文本，不要有任何前言或说明。
"""

# ==========================================
# 主界面布局
# 创建应用的用户界面，包括输入区域和交互元素
# ==========================================
st.markdown("<h1>个人陈述修改</h1>", unsafe_allow_html=True)

# 区域1: 原始文书输入区
with st.expander("**1. 原始文书**", expanded=True):
    # 上传文件区域 - 放在上面
    st.file_uploader("上传文件", type=['docx', 'pdf', 'txt'], key="uploader_ps", 
                     on_change=lambda: st.session_state.update({'ps_content': extract_text_from_file(st.session_state.uploader_ps)}))
    
    # 文本输入区 - 放在下面
    st.text_area(label="", 
                 placeholder="或直接将文本内容复制黏贴在此处",
                 height=150, 
                 key="ps_content")

# 区域2: 新项目信息输入区
with st.expander("**2. 新项目信息**", expanded=True):
    c1, c2 = st.columns(2)
    with c1:
        # 目标学校输入
        target_school = st.text_input("目标学校", placeholder="e.g., Columbia University")
    with c2:
        # 目标专业输入
        target_major = st.text_input("目标专业", placeholder="e.g., MS in Biostatistics")
    
    st.markdown("---")
    # 课程大纲上传
    st.file_uploader("上传课程大纲", type=['docx', 'pdf', 'txt'], key="uploader_curr",
                     on_change=lambda: st.session_state.update({'curr_content': extract_text_from_file(st.session_state.uploader_curr)}))

    # 图片上传区，支持多个图片
    uploaded_images = st.file_uploader("上传图片", type=['png', 'jpg', 'jpeg', 'webp'], accept_multiple_files=True)

    # 课程文本输入
    st.text_area("课程文本:", height=150, key="curr_content")

    st.markdown("---")
    # 写作策略输入区
    st.text_area("3. 写作思路与策略 (可选):", height=100, key="strategy_content", 
                placeholder="例如：这段经历请帮我保留，但要强调我的领导力...")

# ==========================================
# 核心执行逻辑
# 处理用户输入并生成初始文本
# ==========================================
st.divider()
# 开始生成按钮
generate_btn = st.button("1. 开始生成", type="primary")

if generate_btn:
    # 获取用户输入的内容
    final_old_ps = st.session_state.ps_content
    final_new_curr = st.session_state.curr_content
    final_strategy = st.session_state.strategy_content
    
    # 验证必要的输入是否完整
    if not api_key or not final_old_ps.strip() or not target_school:
        st.error("请检查 API Key、旧 PS 内容和目标学校是否完整")
    else:
        # 重置所有状态变量，准备新的生成
        st.session_state['full_response'] = ""
        st.session_state['sections_data'] = [] 
        st.session_state['translation_results'] = {}
        st.session_state['edited_translations'] = {}
        st.session_state['refine_results'] = {}
        st.session_state['preview_results'] = {}
        st.session_state['generation_complete'] = False
        st.session_state['show_sections'] = False
        st.session_state['annotation_processing'] = {}
        st.session_state['annotation_results'] = {}
        st.session_state['original_texts'] = {}
        st.session_state['final_preview_text'] = ""  # 重置最终预览文本
        st.session_state['final_preview_text_cleaned'] = ""  # 重置清理后的预览文本
        st.session_state['confirmed_paragraphs'] = set()  # 重置已确认段落
        st.session_state['confirmed_contents'] = {}  # 重置已确认内容
        
        # 创建一个空白占位符用于显示生成进度
        output_placeholder = st.empty()
        
        with st.spinner(f"正在连接 {model_name} 进行全篇结构分析..."):
            try:
                # 检查是否上传了图片
                has_imgs = True if uploaded_images else False
                # 构建分析提示词
                prompt_text = build_analysis_prompt(target_school, target_major, final_old_ps, final_new_curr, has_imgs, final_strategy)
                
                # 准备内容部分，包括提示词和图片(如果有)
                content_parts = [prompt_text]
                if uploaded_images:
                    for img_file in uploaded_images:
                        content_parts.append(Image.open(img_file))
                
                # 初始化Gemini模型
                model = genai.GenerativeModel(model_name)
                
                # 设置安全过滤级别
                safety_settings = {
                    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
                }

                # 流式生成内容
                response_stream = model.generate_content(
                    content_parts, 
                    stream=True,
                    safety_settings=safety_settings 
                )
                
                # 实时显示生成的内容 - 批处理优化版本
                full_response = ""
                BUFFER_SIZE = 200  # 字符阈值
                UPDATE_INTERVAL = 0.05  # 50ms
                buffer = ""
                last_update = time.perf_counter()

                for chunk in response_stream:
                    try:
                        if chunk.text:
                            buffer += chunk.text  # 暂不清理
                            current_time = time.perf_counter()

                            # 达到阈值或时间间隔时更新
                            if len(buffer) >= BUFFER_SIZE or (current_time - last_update) >= UPDATE_INTERVAL:
                                clean_buffer = clean_asterisks(buffer)
                                full_response += clean_buffer
                                output_placeholder.markdown(full_response + '<span class="streaming-cursor"></span>', unsafe_allow_html=True)
                                buffer = ""
                                last_update = current_time
                    except Exception:
                        pass

                # 最后处理剩余缓冲
                if buffer:
                    clean_buffer = clean_asterisks(buffer)
                    full_response += clean_buffer
                    output_placeholder.markdown(full_response + '<span class="streaming-cursor"></span>', unsafe_allow_html=True)
                
                # 清理和过滤最终响应
                full_response = clean_asterisks(full_response)
                full_response = filter_ai_greeting(full_response)
                output_placeholder.markdown(full_response)
                
                # 保存完整响应
                st.session_state['full_response'] = full_response
                st.session_state['generation_complete'] = True
                
                # 解析响应数据为结构化段落
                raw_sections = full_response.split('===SECTION===')
                parsed_data = []
                
                for sec in raw_sections:
                    if not sec.strip(): continue
                    # 过滤不包含核心标记的段落
                    if "[[LOGIC]]" not in sec and "[[DRAFT]]" not in sec:
                        continue
                        
                    logic_part = ""
                    draft_part = ""
                    if "[[LOGIC]]" in sec:
                        parts = sec.split("[[DRAFT]]")
                        logic_part = parts[0].replace("[[LOGIC]]", "").replace("Part 1:", "").strip()
                        if len(parts) > 1:
                            draft_part = parts[1].replace("Part 2:", "").strip()
                    else:
                        draft_part = sec.strip()
                        
                    parsed_data.append({"logic": logic_part, "draft": draft_part})
                
                # 保存解析后的段落数据
                st.session_state['sections_data'] = parsed_data
                
            except Exception as e:
                st.error(f"生成失败: {e}")

# 显示生成完成的全文
if st.session_state['generation_complete'] and not st.session_state['show_sections']:
    st.markdown("### 生成完成")
    st.markdown(st.session_state['full_response'])
    
    # 显示"开始编辑"按钮
    if st.button("2. 开始编辑段落", key="start_editing_btn", type="primary"):
        st.session_state['show_sections'] = True
        st.rerun()

# ==========================================
# 全篇交互编辑区域
# 提供段落级别的编辑、翻译和修改功能
# ==========================================
if st.session_state['show_sections'] and st.session_state['sections_data']:
    st.divider()
    st.subheader("全篇编辑模式")
    st.caption("请在左侧文本框中直接编辑，或在 `【】` 或 `[]` 中输入修改指令，然后点击下方按钮执行修改。")

    # 使用全局安全设置 safety_settings_interactive

    # 遍历所有段落，为每个段落创建编辑界面
    for i, section_data in enumerate(st.session_state['sections_data']):
        # 在段落标题旁显示状态
        topic = extract_paragraph_topic(section_data['logic'])
        if i in st.session_state['confirmed_paragraphs']:
            st.markdown(f"### {topic} ✅")
        else:
            st.markdown(f"### {topic}")
        
        # 布局：左侧编辑区，右侧逻辑说明
        col_draft, col_logic = st.columns([0.65, 0.35], gap="large")
        
        # 右侧：显示AI修改思路和批注指南
        with col_logic:
            st.info(f"**AI 修改思路 (Logic):**\n\n{section_data['logic']}")
            if "**" in section_data['draft']:
                st.success("已包含高亮修改")
                
            # 添加批注使用指南
            st.markdown("""
            **批注指南:**
            1. 在文本框中使用【】或[]添加批注
            2. 例如：【把这段语气改得更自信】
            3. 点击"执行批注修改"按钮应用修改
            """)
            
            # 检查当前文本是否包含批注，如有则提示用户
            current_text = st.session_state['sections_data'][i]['draft']
            if contains_annotation(current_text):
                st.warning("检测到批注，请点击'执行批注修改'按钮应用修改")

        # 左侧：文本编辑区域
        with col_draft:
            # 检查是否有之前的修改结果，如有则优先显示
            draft_key = f"para_{i}"
            display_text = st.session_state['refine_results'].get(draft_key, section_data['draft'])
            
            # 文本编辑框
            current_draft = st.text_area(
                label="内容编辑",
                value=display_text,
                height=300,
                key=f"draft_p_{i}",
                label_visibility="collapsed"
            )
            
            # 实时保存用户编辑的内容
            st.session_state['sections_data'][i]['draft'] = current_draft
            
            # 检查文本是否包含中文，用于决定输出语言
            has_chinese = contains_chinese(current_draft)
            
            # 操作按钮行
            c_btn1, c_btn2, c_btn3, c_btn4 = st.columns([1, 1, 1, 1])
            
            # 批注修改按钮 - 修改为直接替换原文本并显示预览
            with c_btn1:
                if st.button("执行修改", key=f"btn_refine_{i}"):
                    # 检查是否包含批注标记
                    if contains_annotation(current_draft):
                        with st.spinner("正在根据您的批注优化..."):
                            try:
                                # 保存原始文本用于比较
                                st.session_state['original_texts'][f"para_{i}"] = current_draft
                                
                                # 初始化模型并生成修改后的内容
                                refine_model = genai.GenerativeModel(model_name)
                                res = refine_model.generate_content(
                                    build_refine_prompt(current_draft, has_chinese),
                                    safety_settings=safety_settings_interactive
                                )
                                
                                # 获取优化后的文本
                                refined_text = res.text
                                
                                # 更新会话状态 - 保存修改结果但不直接替换
                                st.session_state['refine_results'][f"para_{i}"] = refined_text
                                st.session_state['annotation_results'][f"para_{i}"] = refined_text
                                
                                # 清除该段落的翻译相关结果
                                if f"trans_{i}" in st.session_state['translation_results']:
                                    del st.session_state['translation_results'][f"trans_{i}"]
                                if f"trans_{i}" in st.session_state['edited_translations']:
                                    del st.session_state['edited_translations'][f"trans_{i}"]
                                if f"preview_trans_{i}" in st.session_state['preview_results']:
                                    del st.session_state['preview_results'][f"preview_trans_{i}"]
                                
                                # 设置批注处理状态
                                st.session_state['annotation_processing'][f"para_{i}"] = True
                                
                                # 显示成功消息并刷新页面
                                st.success("批注修改已应用")
                                st.rerun()
                            except Exception as e:
                                st.error(f"修改失败: {e}")
                    else:
                        st.warning("未检测到批注标记。请在文本中添加【】或[]形式的批注。")

            # 美式英语翻译按钮
            with c_btn2:
                if st.button("🇺🇸翻译", key=f"btn_us_{i}"):
                    with st.spinner("Translating to US English..."):
                        try:
                            # 初始化模型并生成翻译
                            trans_model = genai.GenerativeModel(model_name)
                            res = trans_model.generate_content(
                                build_translate_prompt(current_draft, "US"),
                                safety_settings=safety_settings_interactive
                            )
                            # 保存翻译结果
                            st.session_state['translation_results'][f"trans_{i}"] = {
                                "text": res.text,
                                "style": "US"
                            }
                            # 初始化编辑版本
                            if f"trans_{i}" not in st.session_state['edited_translations']:
                                st.session_state['edited_translations'][f"trans_{i}"] = res.text
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))
            
            # 英式英语翻译按钮
            with c_btn3:
                if st.button("🇬🇧翻译", key=f"btn_uk_{i}"):
                    with st.spinner("Translating to UK English..."):
                        try:
                            # 初始化模型并生成翻译
                            trans_model = genai.GenerativeModel(model_name)
                            res = trans_model.generate_content(
                                build_translate_prompt(current_draft, "UK"),
                                safety_settings=safety_settings_interactive
                            )
                            # 保存翻译结果
                            st.session_state['translation_results'][f"trans_{i}"] = {
                                "text": res.text,
                                "style": "UK"
                            }
                            # 初始化编辑版本
                            if f"trans_{i}" not in st.session_state['edited_translations']:
                                st.session_state['edited_translations'][f"trans_{i}"] = res.text
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))
            
            # 添加确认内容按钮
            with c_btn4:
                # 如果段落尚未确认，显示确认按钮
                if i not in st.session_state['confirmed_paragraphs']:
                    if st.button("✅ 确认内容", key=f"confirm_p_{i}"):
                        logger.info(f"=== 点击确认段落 {i} ===")
                        logger.info(f"current_draft长度: {len(current_draft) if current_draft else 0}")

                        # 调试输出
                        if DEBUG_MODE:
                            st.write(f"调试: 点击确认段落 {i}")
                            st.write(f"调试: confirmed_paragraphs = {st.session_state['confirmed_paragraphs']}")

                        # 标记段落为已确认
                        st.session_state['confirmed_paragraphs'].add(i)
                        logger.info(f"段落 {i} 添加到 confirmed_paragraphs")

                        # 保存当前段落内容到confirmed_contents
                        # 优先从文本框session state获取最新内容
                        textarea_key = f"draft_p_{i}"
                        latest_content = st.session_state.get(textarea_key, current_draft)
                        # 如果latest_content为空或只有空白字符，使用段落原始内容
                        if not latest_content or not latest_content.strip():
                            latest_content = st.session_state['sections_data'][i]['draft']
                            logger.info(f"段落 {i} latest_content为空，使用原始段落内容，长度: {len(latest_content) if latest_content else 0}")
                        st.session_state['confirmed_contents'][i] = latest_content
                        logger.info(f"段落 {i} 保存到 confirmed_contents, 长度: {len(latest_content) if latest_content else 0}")
                        logger.debug(f"段落 {i} 内容前100字符: {latest_content[:100] if latest_content else '空'}")

                        if DEBUG_MODE:
                            st.write(f"调试: confirmed_contents[{i}] = {st.session_state['confirmed_contents'].get(i, 'NOT FOUND')}")

                        # 重建最终预览文本
                        logger.info("开始调用 rebuild_final_preview()")
                        rebuilt_text = rebuild_final_preview()
                        logger.info(f"rebuild_final_preview() 返回长度: {len(rebuilt_text)}")

                        if rebuilt_text and rebuilt_text.strip():
                            st.session_state['final_preview_text'] = rebuilt_text
                            logger.info(f"final_preview_text 设置为重建结果，长度: {len(rebuilt_text)}")

                            if DEBUG_MODE:
                                st.write(f"调试: rebuild结果长度 = {len(rebuilt_text)}")

                            # 清空清理版本，因为内容已更新
                            st.session_state['final_preview_text_cleaned'] = ''
                            logger.info("已清空 final_preview_text_cleaned")

                            logger.info(f"=== 段落 {i} 确认完成 ===")
                            st.success("内容已添加到最终预览")
                        else:
                            logger.error(f"重建的文本为空！段落 {i} 确认失败")
                            st.error(f"无法重建最终预览文本。请检查段落 {i+1} 是否有内容。")

                        st.rerun()
                else:
                    # 如果段落已确认，显示已确认状态
                    st.success("✓ 已确认")
            
            # 显示批注修改结果（如果有）
            if f"para_{i}" in st.session_state['annotation_results']:
                # 获取原始文本和修改后的文本
                original_text = st.session_state['original_texts'].get(f"para_{i}", "")
                refined_text = st.session_state['annotation_results'][f"para_{i}"]
                
                # 高亮显示差异部分
                highlighted_html = highlight_differences(original_text, refined_text)
                
                # 显示修改结果预览
                st.markdown("**批注修改结果预览:**")
                st.markdown(f"""
                <div class="annotation-result-container">
                    {highlighted_html}
                </div>
                """, unsafe_allow_html=True)
                
                # 修改提示文字
                st.caption("修改后的文本已自动更新到上方文本框。黄色高亮部分为修改内容。如不满意，可直接在上方文本框中继续编辑或在【】内添加新批注。")
            
            # 显示翻译结果（如果有）
            trans_key = f"trans_{i}"
            if trans_key in st.session_state['translation_results']:
                trans_data = st.session_state['translation_results'][trans_key]
                st.markdown(f"**{trans_data['style']}式翻译结果:** (可在下方编辑并添加【】批注)")
                
                # 翻译结果编辑区
                edited_trans = st.text_area(
                    "编辑翻译结果",
                    value=st.session_state['edited_translations'].get(trans_key, trans_data["text"]),
                    height=300,
                    key=f"edit_trans_{i}"
                )
                
                # 保存编辑后的翻译结果
                st.session_state['edited_translations'][trans_key] = edited_trans
                
                # 翻译操作按钮
                col1 = st.columns(1)[0]
                
                # 执行翻译批注修改按钮 - 修改为使用英文精修提示词
                with col1:
                    if st.button("执行翻译批注修改", key=f"refine_trans_{i}"):
                        # 检查是否包含批注标记
                        if contains_annotation(edited_trans):
                            with st.spinner("正在根据您的批注优化翻译..."):
                                try:
                                    # 保存原始翻译文本用于比较
                                    st.session_state['original_texts'][f"trans_{i}"] = edited_trans
                                    
                                    # 初始化模型并生成修改 - 使用英文精修提示词
                                    refine_model = genai.GenerativeModel(model_name)
                                    res = refine_model.generate_content(
                                        build_english_refine_prompt(edited_trans),
                                        safety_settings=safety_settings_interactive
                                    )
                                    # 获取修改后的文本
                                    refined_text = res.text
                                    
                                    # 生成预览HTML并保存
                                    preview_html = generate_preview_html(refined_text)
                                    preview_key = f"preview_trans_{i}"
                                    st.session_state['preview_results'][preview_key] = preview_html
                                    
                                    # 保存修改后的文本
                                    st.session_state['edited_translations'][trans_key] = refined_text
                                    
                                    # 显示成功消息并刷新页面
                                    st.success("翻译批注修改已应用")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"修改失败: {e}")
                        else:
                            st.warning("未检测到批注标记。请在文本中添加【】或[]形式的批注。")
                
                # 显示预览结果（如果有）
                preview_key = f"preview_trans_{i}"
                if preview_key in st.session_state['preview_results']:
                    st.markdown("**翻译修改预览结果:**")
                    # 显示带有高亮的HTML预览
                    preview_html = st.session_state['preview_results'][preview_key]
                    st.markdown(preview_html, unsafe_allow_html=True)
                    
                    # 添加提示文字
                    st.caption("✏️ 修改后的文本已自动更新到上方编辑框。黄色高亮部分为修改内容。如不满意，可直接在上方编辑框中继续编辑或在【】内添加新批注。")
        
        # 段落分割线
        st.divider()

    # ==========================================
    # 最终导出区域
    # 提供文档预览和导出功能
    # ==========================================
    st.subheader("最终导出")
    # 导出选项（单列布局）
    # 是否保留加粗高亮
    keep_highlight = st.checkbox("在 Word 中保留加粗高亮", value=True)

    # 自定义页眉选项
    custom_header = st.text_input("自定义页眉专业名称 (可选)", 
                                 value=target_major if target_major else "",
                                 placeholder="例如: Master of Science in Data Science")
    
    # 全文预览
    st.markdown("### 全文预览")
    
    # 显示已确认段落的数量和总段落数
    confirmed_count = len(st.session_state['confirmed_paragraphs'])
    total_paragraphs = len(st.session_state['sections_data'])
    
    if confirmed_count < total_paragraphs:
        st.warning(f"已确认 {confirmed_count}/{total_paragraphs} 段落")
    else:
        st.success(f"已确认全部 {total_paragraphs} 段落")
    
    # 显示最终预览文本
    # 优先显示清理后的版本，如果存在且不为空的话
    cleaned_text = st.session_state.get('final_preview_text_cleaned', '')
    if cleaned_text and cleaned_text.strip():  # 检查清理版本是否存在且非空
        display_text = cleaned_text
        logger.info(f"使用final_preview_text_cleaned作为显示文本")
    else:
        display_text = st.session_state['final_preview_text']
        logger.info(f"使用final_preview_text作为显示文本")

    logger.info(f"=== 显示最终预览 ===")
    logger.info(f"final_preview_text长度: {len(st.session_state['final_preview_text'])}")
    logger.info(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
    logger.info(f"显示文本长度: {len(display_text)}")
    logger.debug(f"final_preview_text前100字符: {st.session_state['final_preview_text'][:100] if st.session_state['final_preview_text'] else '空'}")
    logger.debug(f"display_text前100字符: {display_text[:100] if display_text else '空'}")

    # 如果文本为空，显示提示信息
    if not display_text.strip() and not st.session_state['final_preview_text'].strip():
        logger.warning("最终预览文本为空")
        st.info("请先在上方段落中点击'✅ 确认内容'按钮，将段落添加到最终预览")

    # 调试：直接显示文本内容，检查是否渲染问题
    if DEBUG_MODE:
        st.write(f"调试: display_text类型: {type(display_text)}")
        st.write(f"调试: display_text长度: {len(display_text) if display_text else 0}")
        # 检查字符串中是否有控制字符
        if display_text:
            import string
            printable = set(string.printable)
            non_printable_count = sum(1 for c in display_text[:500] if c not in printable)
            if non_printable_count > 0:
                st.write(f"调试: 前500字符中有 {non_printable_count} 个不可打印字符")
                # 显示非打印字符的位置
                for i, c in enumerate(display_text[:500]):
                    if c not in printable:
                        st.write(f"调试: 位置 {i} 的字符: repr={repr(c)}, ord={ord(c)}")
            # 显示前200字符
            st.code(display_text[:200] if display_text else '空', language='text')
            # 显示repr，查看隐藏字符
            st.write(f"调试: display_text前200字符的repr: {repr(display_text[:200])}")

    def update_final_preview():
        """更新最终预览文本的回调函数"""
        new_value = st.session_state.get('final_preview_text_display', '')
        logger.info(f"=== 文本区域on_change回调被调用 ===")
        logger.info(f"new_value长度: {len(new_value)}")
        logger.info(f"new_value类型: {type(new_value)}")
        if new_value:
            logger.debug(f"new_value前100字符: {new_value[:100]}")

        # 获取当前的清理版本
        current_cleaned = st.session_state.get('final_preview_text_cleaned', '')
        logger.info(f"当前final_preview_text_cleaned长度: {len(current_cleaned)}")

        # 检查新值是否与清理版本相同（可能是AI处理后的自动更新）
        is_same_as_cleaned = new_value == current_cleaned
        logger.info(f"新值与清理版本相同: {is_same_as_cleaned}")

        original_length = len(st.session_state.get('final_preview_text', ''))
        logger.info(f"原final_preview_text长度: {original_length}")

        if new_value:
            st.session_state['final_preview_text'] = new_value
            logger.info(f"更新final_preview_text，新长度: {len(new_value)}")
        else:
            logger.warning(f"new_value为空或None，不更新final_preview_text")

        # 只有在新值与清理版本不同时才清除清理版本（表示用户手动编辑）
        if not is_same_as_cleaned and current_cleaned:
            st.session_state['final_preview_text_cleaned'] = ''
            logger.info("用户手动编辑文本，已清空final_preview_text_cleaned")
        else:
            logger.info("文本未变化或与清理版本相同，保留final_preview_text_cleaned")

    # 确保display_text是字符串
    text_area_value = str(display_text) if display_text is not None else ""

    logger.info(f"=== 文本区域渲染信息 ===")
    logger.info(f"display_text类型: {type(display_text)}")
    logger.info(f"display_text长度: {len(display_text) if display_text else 0}")
    logger.info(f"text_area_value类型: {type(text_area_value)}")
    logger.info(f"text_area_value长度: {len(text_area_value)}")
    logger.debug(f"text_area_value前200字符: {text_area_value[:200] if text_area_value else '空'}")
    logger.info(f"final_preview_text_cleaned存在: {'final_preview_text_cleaned' in st.session_state}")
    logger.info(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
    logger.info(f"final_preview_text_display session state存在: {'final_preview_text_display' in st.session_state}")

    # 在主界面显示调试信息
    if DEBUG_MODE:
        st.markdown("---")
        st.markdown("**文本区域调试信息:**")
        col_debug1, col_debug2 = st.columns(2)
        with col_debug1:
            st.write(f"display_text长度: {len(display_text) if display_text else 0}")
            st.write(f"text_area_value长度: {len(text_area_value)}")
            st.write(f"final_preview_text长度: {len(st.session_state['final_preview_text'])}")
        with col_debug2:
            st.write(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
            st.write(f"final_preview_text_display session state存在: {'final_preview_text_display' in st.session_state}")

        # 显示text_area_value的前100字符
        if text_area_value and len(text_area_value) > 0:
            st.text_area("text_area_value预览", value=text_area_value[:300], height=150, key="debug_preview", disabled=True)
            # 显示repr版本，查看隐藏字符
            st.write(f"text_area_value前200字符repr: {repr(text_area_value[:200])}")
            # 检查是否只包含空白字符
            if text_area_value.strip() == "":
                st.error("text_area_value只包含空白字符！")
                # 显示字符统计
                import string
                printable = set(string.printable)
                non_printable = sum(1 for c in text_area_value[:500] if c not in printable)
                st.write(f"前500字符中不可打印字符数: {non_printable}")
        else:
            st.warning("text_area_value为空")
        st.markdown("---")

    st.text_area(
        "最终文本预览",
        value=text_area_value,
        height=500,
        key="final_preview_text_display",
        on_change=update_final_preview
    )

    # 调试：检查文本区域渲染后的状态
    if DEBUG_MODE:
        st.write(f"调试: 文本区域渲染后，final_preview_text_display session state长度: {len(st.session_state.get('final_preview_text_display', ''))}")
        st.write(f"调试: 文本区域渲染后，final_preview_text长度: {len(st.session_state.get('final_preview_text', ''))}")
        # 尝试用markdown显示一小段内容，检查是否渲染问题
        if display_text and len(display_text) > 50:
            st.markdown(f"**直接Markdown显示前100字符:** {display_text[:100]}...")

    # 去除AI词汇按钮
    st.divider()
    st.markdown("#### 去除AI写作高频词汇")
    st.caption("点击下方按钮去除文本中的AI写作高频词汇和句式（黑名单）。")

    # 在去除AI词汇按钮前添加调试代码
    if DEBUG_MODE:
        st.write(f"调试: final_preview_text长度 = {len(st.session_state['final_preview_text'])}")
        st.write(f"调试: confirmed_paragraphs = {st.session_state['confirmed_paragraphs']}")
        st.write(f"调试: confirmed_contents keys = {list(st.session_state['confirmed_contents'].keys())}")

    if api_key:
        if st.button("🚫 去除AI词汇并生成最终版", type="secondary", use_container_width=True):
            logger.info(f"=== 点击去除AI词汇按钮 ===")
            logger.info(f"final_preview_text长度: {len(st.session_state['final_preview_text'])}")
            logger.info(f"final_preview_text前200字符: {st.session_state['final_preview_text'][:200] if st.session_state['final_preview_text'] else '空'}")
            logger.info(f"final_preview_text_cleaned长度: {len(st.session_state.get('final_preview_text_cleaned', ''))}")
            logger.info(f"final_preview_text_cleaned前200字符: {st.session_state.get('final_preview_text_cleaned', '')[:200] if st.session_state.get('final_preview_text_cleaned') else '空'}")
            logger.info(f"final_preview_text_display session state存在: {'final_preview_text_display' in st.session_state}")
            if 'final_preview_text_display' in st.session_state:
                logger.info(f"final_preview_text_display长度: {len(st.session_state['final_preview_text_display'])}")
                logger.info(f"final_preview_text_display前200字符: {st.session_state['final_preview_text_display'][:200] if st.session_state['final_preview_text_display'] else '空'}")
            logger.info(f"confirmed_paragraphs: {st.session_state['confirmed_paragraphs']}")
            logger.info(f"confirmed_contents keys: {list(st.session_state['confirmed_contents'].keys())}")

            with st.spinner("正在去除AI写作高频词汇..."):
                try:
                    # 获取当前文本 - 优先使用文本区域的当前内容
                    current_text = st.session_state.get('final_preview_text_display', st.session_state['final_preview_text'])
                    logger.info(f"准备处理的文本来源: {'final_preview_text_display' if 'final_preview_text_display' in st.session_state else 'final_preview_text'}")
                    logger.info(f"准备处理的文本长度: {len(current_text) if current_text else 0}")
                    logger.info(f"准备处理的文本前200字符: {current_text[:200] if current_text else '空'}")

                    if not current_text.strip():
                        logger.warning("最终预览文本为空")
                        st.warning("最终预览文本为空")
                    else:
                        logger.info(f"调用AI模型处理文本，长度: {len(current_text)}")
                        logger.info(f"构建prompt并发送到AI...")
                        # 初始化模型
                        refine_model = genai.GenerativeModel(model_name)
                        res = refine_model.generate_content(
                            build_remove_ai_vocab_prompt(current_text),
                            safety_settings=safety_settings_interactive
                        )
                        # 获取处理后的文本
                        cleaned_text = res.text
                        logger.info(f"AI处理完成，返回文本长度: {len(cleaned_text)}")
                        logger.debug(f"处理后文本前500字符: {cleaned_text[:500] if cleaned_text else '空'}")
                        logger.info(f"原始文本长度: {len(current_text)}，处理后长度: {len(cleaned_text)}")

                        # 更新会话状态 - 保存清理版本，保留原始文本
                        st.session_state['final_preview_text_cleaned'] = cleaned_text
                        logger.info(f"final_preview_text_cleaned 已设置，长度: {len(cleaned_text)}")
                        logger.info(f"session state中 final_preview_text_cleaned 现在: {len(st.session_state.get('final_preview_text_cleaned', ''))} 字符")

                        # 清除文本区域的session state，确保下次渲染使用新值
                        if 'final_preview_text_display' in st.session_state:
                            del st.session_state['final_preview_text_display']
                            logger.info("已清除final_preview_text_display，确保下次渲染使用清理版本")

                        # 注意：文本区域的值会在下一次渲染时通过text_area_value自动更新
                        logger.info(f"清理版本已保存到final_preview_text_cleaned")

                        # 检查清理后的文本是否包含Markdown符号
                        markdown_symbols = ['**', '*', '__', '_', '`', '#', '##', '###', '####']
                        found_symbols = [sym for sym in markdown_symbols if sym in cleaned_text]
                        if found_symbols:
                            logger.warning(f"清理文本中仍然包含Markdown符号: {found_symbols}")

                        # 调试：检查 final_preview_text 是否被意外修改
                        logger.info(f"处理后 final_preview_text长度: {len(st.session_state['final_preview_text'])}")

                        # 立即记录最终显示文本的状态
                        cleaned_text_check = st.session_state.get('final_preview_text_cleaned', '')
                        if cleaned_text_check and cleaned_text_check.strip():
                            display_text_should_be = cleaned_text_check
                            logger.info(f"处理完成后，display_text应该显示final_preview_text_cleaned，长度: {len(display_text_should_be)}")
                        else:
                            display_text_should_be = st.session_state['final_preview_text']
                            logger.info(f"处理完成后，display_text应该显示final_preview_text，长度: {len(display_text_should_be)}")

                        st.success("AI词汇已去除，清理版本已生成并显示在文本区域中！")
                        st.rerun()
                except Exception as e:
                    logger.error(f"去除AI词汇失败: {e}")
                    logger.exception(f"完整异常信息:")
                    st.error(f"处理失败: {e}")
    else:
        st.warning("请先配置API Key以使用此功能")
    
    if HAS_DOCX:
        # 准备导出文本 - 优先使用清理版本
        export_text = st.session_state.get('final_preview_text_cleaned') or st.session_state['final_preview_text']
        if not keep_highlight:
            export_text = remove_markdown_bold(export_text)
        
        # 创建Word文档
        docx_file = create_docx_smart(export_text, custom_header)
        
        # 添加下载按钮
        st.download_button(
            label="下载Word文档",
            data=docx_file,
            file_name=f"Personal_Statement_{target_school.replace(' ', '_') if target_school else 'Final'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
        
